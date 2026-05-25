# -*- coding: utf-8 -*-
from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\博士工作\论文工作\（1）小信号模型\Modularized-Small-Signal-Modeling-of-Grid-Forming-Inverters-main")
EIG = ROOT / "EigenAnalysis"
MODE_DIR = EIG / "Control_Mode_Comparison_Results"
SCAN_DIR = EIG / "Control_Parameter_Scan_Results"
TRAJ_DIR = EIG / "Mode_Trajectory_Results"
OUT = ROOT / "研究报告" / "构网型风电机组机电耦合分析_更新模型与文献对照_20260524.docx"

INK = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x5B, 0x65, 0x70)
TABLE_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
ACCENT_FILL = "E8EEF5"
FONT_LATIN = "Calibri"
FONT_CN = "Microsoft YaHei"


def read_csv(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


baseline = read_csv(MODE_DIR / "baseline_torsional_modes.csv")
deltas = read_csv(MODE_DIR / "causal_baseline_deltas.csv")
scan = read_csv(SCAN_DIR / "parameter_scan_summary.csv")
traj = read_csv(TRAJ_DIR / "tracked_torsional_mode_summary.csv")


def font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT_LATIN
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, value, bold=False, color=None, align=None, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    if align:
        p.alignment = align
    r = p.add_run(str(value))
    font(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margin(cell)


def table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for i, width in enumerate(widths):
        grid.gridCol_lst[i].set(qn("w:w"), str(round(width * 1440)))
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
            tcpr = row.cells[i]._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(round(width * 1440)))


def make_table(doc, headers, rows, widths, small=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        shade(table.rows[0].cells[i], TABLE_FILL)
        set_cell_text(table.rows[0].cells[i], text, bold=True, color=INK,
                      align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5 if small else 9.2)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            numeric = i > 0 and len(str(value)) < 18
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.CENTER if numeric else WD_ALIGN_PARAGRAPH.LEFT,
                          size=8.2 if small else 9.1)
    table_widths(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    return table


def para(doc, text="", style=None, bold=False, color=None, size=None, align=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        font(r, size=size, color=color, bold=bold)
    return p


def rich_para(doc, pieces, after=6):
    p = para(doc, after=after)
    for text, bold, italic in pieces:
        r = p.add_run(text)
        font(r, bold=bold, italic=italic)
    return p


def heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    font(r, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    font(r)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    font(r)
    return p


def callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.rows[0].cells[0].width = Inches(6.5)
    shade(table.cell(0, 0), CALLOUT_FILL)
    cell_margin(table.cell(0, 0), top=130, start=160, bottom=130, end=160)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    font(r, bold=True, color=DARK_BLUE)
    p2 = table.cell(0, 0).add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    font(r2)
    table_widths(table, [6.5])
    para(doc, after=3)


def figure(doc, path, caption, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(caption)
    font(r, size=9, color=GRAY, italic=True)


def pct(a, b):
    return (float(b) - float(a)) / float(a) * 100


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT_LATIN
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, before, after, color in (
    ("Heading 1", 16, 16, 8, BLUE),
    ("Heading 2", 13, 12, 6, BLUE),
    ("Heading 3", 12, 8, 4, DARK_BLUE),
):
    st = styles[name]
    st.font.name = FONT_LATIN
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = sec.header.paragraphs[0]
header.text = ""
header.paragraph_format.space_after = Pt(0)
r = header.add_run("构网型风电机组机电耦合分析 | 更新模型与文献对照")
font(r, size=9, color=GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run("研究工作稿  |  2026 年 5 月 24 日  |  第 ")
font(rf, size=9, color=GRAY)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
footer._p.append(fld)
rf2 = footer.add_run(" 页")
font(rf2, size=9, color=GRAY)

# First-page masthead
para(doc, "技术分析报告", bold=True, size=11, color=BLUE, after=12)
para(doc, "构网型风电机组机电耦合特性分析", bold=True, size=24, color=INK, after=5)
para(doc, "基于更新的小信号模型、四拓扑因果对照与 Zotero 关键文献的重新分析", size=13, color=GRAY, after=18)
meta = [
    ("研究对象", "PMSG 全功率构网型风电机组，两质量块传动链，约 2 Hz 扭振模态"),
    ("模型版本", "含气动转矩线性化、传动链自阻尼、MPPT 斜率反馈与 APCAD 支路的最新程序"),
    ("结果来源", "2026-05-24 重新导出的统一状态空间模型及重新运行的 MATLAB 扫描结果"),
    ("文献依据", "Zotero 文库：机电耦合振荡、构网型控制及 Type-IV GFM-WT 关键论文"),
    ("结论用途", "作为论文机理分析、控制设计与后续 Simulink/RT-LAB 验证的工作底稿"),
]
for label, value in meta:
    p = para(doc, after=3)
    r = p.add_run(f"{label}：")
    font(r, bold=True, color=INK)
    r = p.add_run(value)
    font(r)
para(doc, after=7)
callout(
    doc,
    "结论摘要",
    "新模型表明，单纯将同步方式由 PLL 换为 GFM 并未改变约 2 Hz 扭振阻尼；真正使阻尼降低的是直流电压控制由网侧迁移到机侧后形成的“直流链-电磁转矩-传动链”反馈通道。APCAD 在网侧有功参考中注入针对该频带的阻尼功率，使扭振模态重新左移并提高阻尼比。当前结论可支撑机理论证和控制路线，但数值尚需目标机组参数与非线性/半实物验证确认。"
)

heading(doc, "1 研究问题与文献定位", 1)
para(doc, "本报告重新回答三个直接关系论文立论的问题：构网型风电机组为什么可能增强机电耦合扭振；当前四个对照模型究竟证明了哪一条因果通道；围绕该通道设计的附加阻尼控制是否有效、下一步怎样完成非线性实验闭环。")
para(doc, "在 Zotero 文库中，与当前 PMSG 全功率模型最直接对应的研究形成了清晰证据链。Liu 等在 2024 年关于直流电压控制的研究中，将机侧 DVC 与传动链扭振联系起来，并在小信号推导后用非线性时域仿真验证。另一篇 2024 年研究比较了 GFM-GWT 与 GFM-MWT，指出机侧 DVC 配置会引入不利阻尼，而其对扭振频率的改变相对有限。Udawatte 等在 2026 年提出 APCAD：通过 GSC 有功控制支路施加阻尼，避免直接干扰 MSC 的直流电压调节，并进一步开展实时实验验证。[1]-[3]")
make_table(doc,
           ["文献与模型对应点", "可支撑的研究论断", "在本程序中的落实"],
           [
               ("Liu et al., 2024, TEC [1]", "MSC-DVC 使直流动态与电磁转矩、传动链发生耦合", "GFM-MWT 与 GFM-GWT 的控制位置对照"),
               ("Liu et al., 2024, TSTE [2]", "GFM-MWT 阻尼低于 GFM-GWT；频率变化有限", "四拓扑基准、参与度与模态轨迹"),
               ("Udawatte et al., 2026, TPEL [3]", "GSC-APC 中 BPF+超前+增益可提供附加阻尼", "GFM-MWT+AD / APCAD 与 K_damp 扫描"),
           ],
           [1.78, 2.45, 2.27], small=True)

heading(doc, "2 更新模型：现阶段能表征什么", 1)
heading(doc, "2.1 四类控制拓扑的因果设计", 2)
make_table(doc,
           ["模型", "机侧变流器 MSC", "网侧变流器 GSC", "对照目的"],
           [
               ("GFL-WT", "MPPT", "DVC + PLL", "跟网型基线"),
               ("GFM-GWT", "MPPT", "DVC + GFM", "仅改变同步/成网方式"),
               ("GFM-MWT", "DVC", "GFM", "隔离直流控制位置迁移影响"),
               ("GFM-MWT+AD", "DVC", "GFM + APCAD", "验证附加阻尼作用"),
           ],
           [1.18, 1.45, 1.55, 2.32])
para(doc, "该设计不是简单的“GFL 与 GFM 二选一”比较，而是逐层替换控制环节。因此，只有在前一步物理对象保持一致的前提下，后一步模态变化才可归属于该控制结构改变。")

heading(doc, "2.2 本轮已补入的机械与气动反馈", 2)
para(doc, "本轮模型不再把机械输入转矩仅作为独立外扰。两质量块传动链保留状态 Δωt、Δωg 与 Δθtw，并加入工作点附近的气动转矩线性化：")
para(doc, "ΔTaero = -Daero Δωt + Kv,aero Δvw + Kβ,aero Δβ", bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
para(doc, "其中负斜率项 -Daero Δωt 为 MPPT 区域的气动恢复/阻尼通道；模型还加入 turbine/generator 自阻尼 Dt、Dg，并将风速扰动 Δvw 与桨距扰动 Δβ 作为可扩展输入。GFM-MWT 的有功参考同时引入 k_p_mppt Δωg，使机电功率平衡具有实际转速反馈路径。")
make_table(doc,
           ["模块", "当前包含状态/通道", "研究意义", "目前边界"],
           [
               ("气动侧", "Daero、Kv,aero、Kβ,aero；vw、β 输入", "避免刚体慢模态缺少恢复机制", "MPPT 工作点；桨距增益暂设为 0"),
               ("传动链", "两质量块、轴刚度/阻尼、Dt/Dg", "直接捕捉轴系扭振", "参数为 1 MW 级暂定值"),
               ("电磁与直流链", "PMSG、MSC/GSC、Cdc 与 DVC", "形成直流链至转矩反馈路径", "待真实电机/变流器参数替换"),
               ("阻尼控制", "2 Hz BPF、超前环节、K_damp", "针对目标扭振模态施加电气阻尼", "增益需经非线性限幅验证"),
           ],
           [1.1, 1.83, 1.8, 1.77], small=True)
callout(doc, "建模判断", "当前模型已经能够用于回答“控制结构如何影响目标扭振模态”的小信号机理问题；它尚不是目标实机的最终参数化模型。尤其是气动导数、轴系参数、PMSG 参数、直流电容及控制限幅必须在论文最终定量结论和实时实验前落实。")

heading(doc, "3 分析方法：不只看参与度", 1)
para(doc, "本轮分析将文献中的典型思路转化为三层证据。参与因子用于识别某一特征值是否确为轴系模态；极点实部与阻尼比用于判断控制结构对该模态是增阻尼还是减阻尼；参数扫描与连续模态轨迹用于判断结论是否只在一个工作点偶然成立。")
make_table(doc,
           ["层次", "所用量", "判断问题", "本轮实现"],
           [
               ("模态身份", "参与因子", "约 2 Hz 极点是否由传动链主导", "θtw、ωg、ωt 参与度排序"),
               ("局部阻尼", "λ = σ ± jω；ζ = -σ/|λ|", "结构改变是否使扭振模态左右移动", "基准四拓扑对照"),
               ("趋势与鲁棒性", "特征向量相关连续跟踪", "参数变化下是否仍跟踪同一模态", "h、mp、kpdc、kidc、K_damp 轨迹"),
               ("系统边界", "max Re(λ)", "目标模态改善时全系统是否稳定", "二维稳定区域与增益扫描"),
           ],
           [1.0, 1.65, 2.22, 1.63], small=True)
para(doc, "因此，在论文中不应写成“参数影响仅由参与度证明”。更严谨的逻辑是：先用参与因子确定扭振模态，再用极点轨迹和阻尼变化说明参数作用方向，最后用最大实部检查其他模态是否成为新的稳定性约束。")

heading(doc, "4 基准结果：两质量块扭振模态的控制结构影响", 1)
rows = []
for r in baseline:
    rows.append((
        r["Model"],
        f'{float(r["FrequencyHz"]):.4f}',
        f'{float(r["Sigma"]):.5f}',
        f'{float(r["DampingRatio"]):.5f}',
        f'{float(r["MaxReal"]):.5f}',
        "稳定" if r["Stable"] in ("1", "true", "True") else "不稳定",
    ))
make_table(doc,
           ["模型", "扭振频率/Hz", "σ / s⁻¹", "阻尼比 ζ", "最大实部/s⁻¹", "整体状态"],
           rows, [1.10, 1.02, 1.06, 1.04, 1.28, 1.00], small=True)
para(doc, "在更新模型下，四类拓扑的最大特征值实部均为负值，说明气动恢复路径和传动链自阻尼补入后，基准点不存在此前由机械慢模态缺失反馈造成的右半平面极点。与此同时，所有拓扑的目标模态频率仍集中在约 2 Hz，说明当前观察对象确实是同一传动链扭振模态。")
figure(doc, MODE_DIR / "torsional_damping_comparison.png",
       "图 1  四类控制拓扑下约 2 Hz 扭振模态阻尼比对比（本轮 MATLAB 重新计算结果）")

heading(doc, "5 因果解释：构网型何时增加机电耦合风险", 1)
delta_rows = []
factor_cn = {
    "Synchronization effect (PLL to GFM)": "同步方式变化：PLL → GFM",
    "DC-link allocation effect (GSC-DVC to MSC-DVC)": "直流控制迁移：GSC-DVC → MSC-DVC",
    "APCAD damping effect": "附加阻尼：加入 APCAD",
}
for r in deltas:
    delta_rows.append((
        factor_cn[r["Factor"]],
        f'{r["FromModel"]} → {r["ToModel"]}',
        f'{float(r["DeltaDampingRatio"]):+.6f}',
        f'{float(r["DeltaSigma"]):+.6f}',
    ))
make_table(doc, ["逐级变化", "拓扑转换", "Δζ", "Δσ / s⁻¹"], delta_rows, [2.25, 1.65, 1.28, 1.32])
para(doc, "第一步，`GFL-WT → GFM-GWT` 的阻尼变化为数值零量级。这意味着在当前配置中，只把网侧同步与成网机制由 PLL 改为 GFM，而保留 MSC-MPPT 与网侧直流平衡，不足以证明构网控制本身必然激发轴系扭振。")
para(doc, "第二步，`GFM-GWT → GFM-MWT` 后，阻尼比由 0.012227 降至 0.006617，下降约 45.88%，极点实部右移 0.070406 s⁻¹，但仍留在左半平面。这一步只改变了直流电压控制承担位置：MSC-DVC 将直流电压波动直接转化为 PMSG 电磁转矩调整，再反馈到两质量块传动链。该闭环通道就是当前模型中“构网型增加机电耦合振荡可能性”的具体机制，与文献 [1]-[2] 的结论一致。")
para(doc, "第三步，`GFM-MWT → GFM-MWT+AD` 后，阻尼比从 0.006617 提升至 0.015031，提高约 127.15%，扭振极点左移 0.104538 s⁻¹。APCAD 状态在扭振模态参与因子中进入前列，说明控制器并非仅改变某个不相关的电气模态，而是实际作用于目标轴系振荡。")
figure(doc, MODE_DIR / "torsional_real_part_comparison.png",
       "图 2  分阶段拓扑替换下扭振特征值实部对比；实部越负，目标扭振衰减越快")
callout(doc, "可用于论文的关键表述", "本研究结果不支持“采用 GFM 同步机制即必然恶化轴系扭振”的宽泛结论。结果表明，构网型全功率风机中更关键的风险来源是为获得独立直流电压支撑而采用 MSC-DVC 后形成的直流链能量调节至电磁转矩的反馈通道；该通道降低了约 2 Hz 扭振模态阻尼，而在 GSC 有功参考端配置的 APCAD 能够为同一模态提供有效附加阻尼。")

heading(doc, "6 参与因子与模态归属", 1)
para(doc, "本轮参与因子结果显示，GFM-MWT 的目标模态主要由 `theta_tw`（约 0.499）、`omega_g`（约 0.444）和 `omega_t`（约 0.055）主导；这些是典型两质量块扭振状态。加入 APCAD 后，这三个机械状态仍占主导，同时 `x_bp1` 与 `x_bp2` 分别进入约 0.025 的参与度，说明带通提取支路与原扭振模态发生了预期的控制耦合。")
make_table(doc,
           ["模型", "主要参与状态", "物理解释"],
           [
               ("GFM-MWT", "θtw: 0.499；ωg: 0.444；ωt: 0.055", "传动链扭角与两端转速主导的轴系模态"),
               ("GFM-MWT+AD", "θtw: 0.464；ωg: 0.407；ωt: 0.052", "原轴系模态仍被跟踪，身份未改变"),
               ("GFM-MWT+AD", "xbp1: 0.025；xbp2: 0.025；xlead: 0.004", "APCAD 进入目标模态并施加阻尼"),
           ],
           [1.28, 2.62, 2.60])
para(doc, "但参与度本身不能告诉我们调大参数会使极点向左还是向右移动。因此，下一节的连续轨迹分析是参数研究的主体，而参与因子是模态识别依据。")

heading(doc, "7 控制参数扫描与连续轨迹结果", 1)
heading(doc, "7.1 二维稳定域扫描", 2)
summary_rows = []
scan_names = {"h_mp": "h / mp", "dvc": "kpdc / kidc", "rpc": "kpq / kiq", "K_damp": "K_damp"}
for r in scan:
    summary_rows.append((
        r["Model"],
        scan_names.get(r["Scan"], r["Scan"]),
        f'{r["StablePoints"]}/{r["Points"]}',
        f'{float(r["MinDampingRatio"]):.5f} ~ {float(r["MaxDampingRatio"]):.5f}',
        f'{float(r["RecommendedDampingRatio"]):.5f}',
    ))
make_table(doc,
           ["模型", "扫描参数", "稳定点", "扭振阻尼比范围", "稳定域内推荐 ζ"],
           summary_rows, [1.18, 1.18, 0.93, 1.80, 1.41], small=True)
para(doc, "与补充机械闭环前的扫描不同，本轮所有参数组都出现了稳定区域。`kpdc/kidc` 对 GFM-MWT 的扭振阻尼调节最直接，符合 MSC-DVC 位于电磁转矩入口的物理结构；`kpq/kiq` 对目标扭振阻尼影响较弱，主要作为全系统电气稳定性的约束；`h/mp` 与 APCAD 通道存在交互，因而在加入附加阻尼后需要联合整定。")
figure(doc, SCAN_DIR / "dvc_stability_maps.png",
       "图 3  MSC 直流电压环参数扫描：稳定域与目标扭振阻尼比分布")

heading(doc, "7.2 为什么必须看极点移动轨迹", 2)
traj_rows = []
for r in traj:
    if r["Parameter"] in ("h", "mp", "k_pdc", "k_idc", "K_damp"):
        traj_rows.append((
            r["Model"],
            r["Parameter"],
            f'{float(r["min_DampingRatio"]):.5f}',
            f'{float(r["max_DampingRatio"]):.5f}',
            f'{float(r["min_MaxReal"]):.5f}',
            f'{float(r["max_MaxReal"]):.5f}',
        ))
make_table(doc,
           ["模型", "参数", "ζ 最小", "ζ 最大", "maxRe 最小", "maxRe 最大"],
           traj_rows, [1.22, 0.95, 1.03, 1.03, 1.13, 1.14], small=True)
para(doc, "连续轨迹算法按特征向量相关性跟踪同一对扭振极点，避免在参数变化时误把附近其他模态当成目标模态。结果显示：未加阻尼的 GFM-MWT 对 `h`、`mp` 的扭振变化相对温和；机侧 DVC 参数会改变该通道的阻尼贡献并可能引发其他不稳定模态；在 APCAD 模型中，`K_damp` 可以将阻尼比提高至约 0.03365，但增益方向反转或过度偏移会使目标模态变为负阻尼。")
figure(doc, TRAJ_DIR / "K_damp_tracked_trajectory.png",
       "图 4  APCAD 增益变化下经连续跟踪的扭振模态轨迹及稳定性边界")

heading(doc, "8 阻尼控制设计解释与下一轮整定", 1)
para(doc, "当前 APCAD 的信号路径与文献 [3] 一致：以发电机转速偏差为测量量，经中心频率约 `2 Hz` 的带通滤波器提取轴系振荡分量，再通过超前环节补偿从网侧功率调制到机侧电磁转矩响应之间的相位滞后，最终将小幅阻尼功率叠加到 VSG 有功参考。")
make_table(doc,
           ["控制环节", "当前程序设置", "作用", "必须补做的验证"],
           [
               ("带通滤波 BPF", "fdamp = 2.0 Hz；ζdamp = 0.20", "只提取目标轴系频带，抑制慢速功率变化与高频噪声", "频率漂移及噪声敏感性"),
               ("超前补偿", "Tlead = 1/ωdamp；α = 0.30", "在目标频率对齐阻尼转矩相位", "SCR、Cdc、风速变化下相位裕度"),
               ("阻尼增益", "Kdamp = -5.0×10⁶", "决定附加阻尼强度与极点左移量", "限幅、电流峰值、直流电压峰值"),
               ("注入位置", "GSC-VSG 有功参考端", "通过直流链间接形成电磁阻尼，不直接扰乱 MSC-DVC", "与频率支撑/限功率功能协调"),
           ],
           [1.18, 1.62, 1.93, 1.77], small=True)
para(doc, "基准增益 `Kdamp = -5.0×10^6` 已把阻尼比提高到 `0.015031`；扫描中稳定点的最大目标阻尼出现在 `Kdamp = -1.5×10^7` 左右，对应 `ζ ≈ 0.033652`。该数值只能作为非线性试验候选上界，而不能直接作为最终设置，因为小信号模型尚未包含功率饱和、电流限幅、PWM 非线性、测量延迟与故障穿越逻辑。")
figure(doc, SCAN_DIR / "kdamp_scan.png",
       "图 5  APCAD 增益扫描：目标阻尼提升与全系统稳定边界的共同约束")

heading(doc, "9 非线性验证与 RT-LAB 实验路线", 1)
para(doc, "小信号分析回答的是工作点附近的极点位置与阻尼机理；要形成完整的构网型风机机电耦合实验链，必须证明同一控制作用在非线性扰动、限幅以及实时计算条件下仍可重复出现。建议后续验证按以下顺序实施。")
make_table(doc,
           ["阶段", "软件/平台", "模型与激励", "输出与判据"],
           [
               ("A. 参数落地", "MATLAB 脚本", "以目标 PMSG/轴系/变流器参数替换 TEMP 值，重跑特征值", "2 Hz 模态身份、阻尼比、稳定域可复现"),
               ("B. EMT 非线性验证", "Simulink / Simscape Electrical", "GFM-GWT、GFM-MWT、GFM-MWT+AD；相角阶跃、电压跌落、风速阶跃", "ωt/ωg/轴扭矩/有功/vdc 包络衰减与峰值"),
               ("C. 鲁棒性验证", "Simulink 批处理", "SCR、Cdc、风速、h/mp、DVC 与 Kdamp 联合变化", "阻尼控制不以其他模态失稳或过流为代价"),
               ("D. 实时半实物", "RT-LAB / OPAL-RT", "功率级与电机/风轮实时被控对象；控制器实时执行", "采样延迟、量化、限幅条件下衰减时间与稳定性"),
           ],
           [1.12, 1.35, 2.35, 1.68], small=True)
heading(doc, "9.1 最小可执行非线性对照试验", 2)
number(doc, "在同一额定功率、风速、SCR 与直流电容下，分别运行 `GFM-GWT`、`GFM-MWT`、`GFM-MWT+AD`，保持扰动幅值一致。")
number(doc, "施加网侧相角阶跃（建议从小扰动开始，例如 1°/2°）和电压跌落，记录 `omega_t`、`omega_g`、`theta_tw` 或轴转矩、`P_g` 与 `v_dc`。")
number(doc, "对约 2 Hz 分量进行包络拟合或 Prony/矩阵铅笔识别，比较衰减率与阻尼；同时统计直流电压及电流峰值。")
number(doc, "将 `Kdamp = -5.0×10^6` 作为基线控制值，以 `-1.5×10^7` 作为候选增强值，在限幅开启状态下逐步验证，拒绝仅凭小信号最大阻尼点直接上机。")
heading(doc, "9.2 RT-LAB 环节需要特别验证的量", 2)
bullet(doc, "实时步长、采样频率与控制延迟是否使 BPF/超前补偿的目标相位发生偏移。")
bullet(doc, "APCAD 输出限幅后，扭振衰减是否仍优于未加阻尼模型，以及是否造成有功指令或直流电压冲击。")
bullet(doc, "在弱网、风速变化和不同直流电容下，目标模态是否仍保持在控制器带通有效频带内。")
bullet(doc, "实验波形中识别出的频率与小信号预测约 `2 Hz` 是否一致；不一致时应先更新轴系/气动参数，再重新整定滤波器。")

heading(doc, "10 结论、边界与论文写作版本", 1)
heading(doc, "10.1 当前可以成立的结论", 2)
number(doc, "经气动与机械反馈闭环补充后，四拓扑基准运行点均稳定，说明本轮模型已消除原先由机械恢复机制不完整造成的虚假慢不稳定结论。")
number(doc, "目标机电耦合模态位于约 `2 Hz` 且由轴扭角和两端转速主导，属于两质量块传动链扭振模态。")
number(doc, "控制风险不能笼统归于 GFM 同步机制；在当前模型中，MSC 承担 DVC 后形成的转矩反馈通道使目标扭振阻尼显著降低。")
number(doc, "APCAD 的带通-超前-增益支路能让同一扭振模态明显左移，在基准点把阻尼比从 `0.006617` 提升至 `0.015031`。")
number(doc, "传统参数分析不应止于参与因子；连续极点轨迹和全系统最大实部是阐明趋势与整定边界的必要证据。")
heading(doc, "10.2 不能提前宣称的内容", 2)
bullet(doc, "不能将当前阻尼数值直接等同于某一真实兆瓦级机组的最终指标，因为关键机电参数仍为暂定数据。")
bullet(doc, "不能仅凭小信号极点宣称控制器已满足大扰动、限幅和故障穿越要求；这些必须由 EMT 与 RT-LAB 验证。")
bullet(doc, "不能宣称所有 GFM 架构必然比 GFL 更容易扭振；当前证据定位的是特定直流电压控制配置形成的耦合通道。")
callout(doc, "可直接纳入论文结果分析的段落", "在含气动转矩线性化、两质量块传动链、PMSG、直流链与全功率变流控制的统一小信号模型中，约 2 Hz 扭振模态主要由轴扭角及机电两端转速状态参与。保持 MSC-MPPT 时，将网侧同步机制由 PLL 替换为 GFM 未显著改变该模态阻尼；当直流电压调节迁移至机侧并经电磁转矩作用于传动链时，扭振阻尼比由 0.012227 降至 0.006617，表明机侧直流电压控制引入了不利机电反馈。针对该通道，在 GSC 有功控制端配置 APCAD 后，扭振阻尼比提高至 0.015031，且附加阻尼状态进入目标模态参与因子排序，证明该控制能够针对性提高轴系模态衰减能力。")

heading(doc, "参考文献", 1)
refs = [
    "[1] S. Liu, H. Wu, T. Bosma, and X. Wang, “Impact of DC-Link Voltage Control on Torsional Vibrations in Grid-Forming PMSG Wind Turbines,” IEEE Transactions on Energy Conversion, vol. 39, no. 4, pp. 2631–2644, 2024, doi: 10.1109/TEC.2024.3394753.",
    "[2] S. Liu, R. G. Cirstea, H. Wu, T. Bosma, and X. Wang, “Comparative Evaluation of Converter Control Impact on Torsional Dynamics of Type-IV Grid-Forming Wind Turbines,” IEEE Transactions on Sustainable Energy, vol. 15, no. 4, pp. 2803–2816, 2024, doi: 10.1109/TSTE.2024.3444474.",
    "[3] H. E. Udawatte, A. Asbafkan, M. H. Ravanji, and B. Bahrani, “Active Power Control-Based Damping for Torsional Oscillations in Grid-Forming Type-IV Wind Turbines,” IEEE Transactions on Power Electronics, vol. 41, no. 6, pp. 9940–9955, 2026, doi: 10.1109/TPEL.2026.3653039.",
    "[4] 马亦卓，谭贝斯，徐晋，等，“基于虚拟同步控制的构网型直驱风机轴系扭振特性分析,” 2025。来源：Zotero 构网型控制文献集合。",
    "[5] 李海盼，年珩，胡彬，李萌，“风电并网系统宽频振荡分析与抑制方法综述,” 2023。来源：Zotero 文库。",
]
for ref in refs:
    p = para(doc, ref, after=5)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)

heading(doc, "附录 A 复现文件与输出位置", 1)
make_table(doc,
           ["用途", "程序/结果文件"],
           [
               ("四拓扑对照分析", r"EigenAnalysis\Compare_Control_Mode_Run.m"),
               ("控制参数稳定域扫描", r"EigenAnalysis\Scan_GFM_Control_Parameters_Run.m"),
               ("扭振模态连续轨迹", r"EigenAnalysis\Track_GFM_Mode_Trajectories.m"),
               ("最新统一模型", r"EigenAnalysis\Unified_WT_PMSG_*.mat"),
               ("基准与因果结果", r"EigenAnalysis\Control_Mode_Comparison_Results"),
               ("参数扫描结果", r"EigenAnalysis\Control_Parameter_Scan_Results"),
               ("模态轨迹结果", r"EigenAnalysis\Mode_Trajectory_Results"),
           ],
           [1.85, 4.65])
para(doc, "注：本报告引用的结果均在 2026 年 5 月 24 日从最新导出的统一模型重新计算生成。MATLAB 启动时出现的旧路径警告不影响本轮状态空间模型求解与结果文件输出。", size=9.5, color=GRAY, after=3)

doc.core_properties.title = "构网型风电机组机电耦合特性分析：更新模型与文献对照"
doc.core_properties.subject = "PMSG GFM-WT torsional oscillation, APCAD and validation plan"
doc.core_properties.author = "研究工作稿"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(str(OUT))
